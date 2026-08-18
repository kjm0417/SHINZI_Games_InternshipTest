from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\GitHub\SHINZI_Games_InternshipTest")
OUT = ROOT / "Docs" / "SHINZI_Project_Handoff_Guide.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
TEXT = "263238"
MUTED = "60717D"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F8FB"
PALE_GREEN = "EAF5EE"
PALE_AMBER = "FFF6DF"
PALE_RED = "FDECEC"
WHITE = "FFFFFF"
GRID = "B7C3CD"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index < len(table.rows) - 1:
                for paragraph in cell.paragraphs:
                    p_pr = paragraph._p.get_or_add_pPr()
                    if p_pr.find(qn("w:keepNext")) is None:
                        p_pr.append(OxmlElement("w:keepNext"))


def set_run_font(run, size=11, bold=False, color=TEXT, italic=False, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def set_cell_text(cell, text, *, bold=False, color=TEXT, size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def set_header(header):
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("SHINZI GAMES INTERNSHIP TEST  |  PROJECT HANDOFF"), size=8.5, bold=True, color=MUTED)


def set_footer(footer):
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    add_page_number(p)


def set_section_header_footer(section):
    # Word/PDF 변환 환경에 따라 짝수 쪽 머리글·바닥글이 별도 스토리로
    # 처리될 수 있으므로 양쪽을 명시적으로 같은 형식으로 구성한다.
    set_header(section.header)
    set_footer(section.footer)
    set_header(section.even_page_header)
    set_footer(section.even_page_footer)


def configure_document(doc):
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    set_section_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    # 본문 전체에 keep-together를 상속시키면 큰 표·구조도 앞의 제목까지
    # 한 덩어리로 묶여 페이지 상단을 넘어갈 수 있으므로 필요한 문단에만 적용한다.
    normal.paragraph_format.keep_together = False

    for style_name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_together = False
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_together = True


def add_para(doc, text="", *, bold_prefix=None, color=TEXT, size=11, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = keep
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), size=size, bold=True, color=color)
        set_run_font(p.add_run(text[len(bold_prefix):]), size=size, color=color, italic=italic)
    else:
        set_run_font(p.add_run(text), size=size, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0, keep_next=False):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = keep_next
    set_run_font(p.add_run(text), size=10.5, color=TEXT)
    return p


def add_number(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    set_run_font(p.add_run(f"{number}.  {text}"), size=10.5, color=TEXT)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.083)
    p.paragraph_format.right_indent = Inches(0.083)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), "8796A3")
        borders.append(border)
    p_pr.append(borders)
    set_run_font(p.add_run(title + "\n"), size=10.5, bold=True, color=accent)
    set_run_font(p.add_run(body), size=10, color=TEXT)


def add_table(doc, headers, rows, widths_dxa, header_fill=LIGHT_BLUE, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, bold=True, color=NAVY, size=9.4, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(hdr.cells[i], header_fill)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_layer_box(doc, title, body, fill):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.083)
    p.paragraph_format.right_indent = Inches(0.083)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), "8796A3")
        borders.append(border)
    p_pr.append(borders)
    set_run_font(p.add_run(title + "\n"), size=10.5, bold=True, color=NAVY)
    set_run_font(p.add_run(body), size=9.5, color=TEXT)


def add_arrow(doc, label="의존·호출"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(f"↓  {label}"), size=9, bold=True, color=BLUE)


def add_cover(doc):
    add_para(doc, "PROJECT HANDOFF GUIDE", size=10, color=BLUE, after=4, keep=True)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run("SHINZI Games\n인턴십 과제 프로젝트"), size=30, bold=True, color=NAVY)
    add_para(doc, "Unity 탑다운 1:1 액션 게임 — 구현 구조·데이터 파이프라인·인계 가이드", size=14, color=MUTED, after=22)

    add_callout(
        doc,
        "문서 기준",
        "2026-08-18 저장소의 실제 코드·ScriptableObject·Addressables 설정·씬 구성을 기준으로 작성했습니다. "
        "현재 구현과 향후 확장안을 분리해 기술하며, 구현되지 않은 기능은 현재 기능처럼 표현하지 않습니다.",
        fill=PALE_BLUE,
    )

    add_table(
        doc,
        ["항목", "내용"],
        [
            ("프로젝트 유형", "Unity / 탑다운 1:1 액션 프로토타입"),
            ("핵심 입력", "WASD 이동, 마우스 에임·공격, 대시 입력"),
            ("핵심 데이터 흐름", "Excel → 검증 → ScriptableObject → 런타임 초기화"),
            ("리소스 로딩", "Addressables 기반 플레이어·AI·무기·픽업·투사체 로드"),
            ("맵 운용", "현재 1개 맵을 씬에 고정 배치"),
        ],
        [2200, 7160],
        font_size=9.7,
    )

    add_para(doc, "이 문서는 인계받는 사람이 게임 흐름과 클래스 책임을 빠르게 파악하고, 데이터 변경·리소스 추가·오류 점검을 재현할 수 있도록 구성했습니다.", size=10.5, color=MUTED, italic=True, before=14, after=0)
    doc.add_page_break()


def add_contents(doc):
    doc.add_heading("문서 구성", level=1)
    add_table(
        doc,
        ["구분", "확인할 내용"],
        [
            ("1. 게임 설명", "플레이 목표, 조작, 승패·진행 규칙"),
            ("2. 전체 설계 구조", "레이어 구조, 런타임 시작·종료 흐름"),
            ("3. 역할 분리", "Manager·Controller·Brain·Component·Data의 책임"),
            ("4. 전투·AI·무기", "공통 캐릭터 기능, AI 판단, 무기 런타임 구조"),
            ("5. 데이터 도구", "Excel→SO 2-Pass 변환, 검증, 참조 관계"),
            ("6. Addressables", "주소 저장, 비동기 생성, 해제·버전 보호"),
            ("7. 확장·인계", "현재 제약, 확장 지점, 제출 전 점검"),
        ],
        [1600, 7760],
    )
    add_callout(
        doc,
        "표기 원칙",
        "[현재]는 저장소에 구현된 동작, [확장]은 요구가 생겼을 때 적용할 설계 방향, [점검]은 제출·인계 전에 확인해야 할 항목을 뜻합니다.",
        fill=PALE_GREEN,
        accent=DARK_BLUE,
    )


def add_game_overview(doc):
    doc.add_heading("1. 게임 설명", level=1)
    add_para(doc, "플레이어가 제한 시간 동안 한 명의 AI와 전투하는 탑다운 1:1 액션 게임입니다. 맵에 생성되는 무기를 획득해 교체할 수 있고, 승리 횟수에 따라 다음 매치에서 사용할 AI와 드롭 구성이 달라집니다.")

    doc.add_heading("핵심 플레이", level=2)
    add_bullet(doc, "고정 카메라 환경에서 WASD로 이동하고 마우스 위치를 향해 회전·조준합니다.")
    add_bullet(doc, "공격과 대시를 사용할 수 있으며, 공격·대시 쿨다운은 HUD에 표시됩니다.")
    add_bullet(doc, "맵에 생성된 다른 무기를 획득하면 현재 무기 데이터와 손에 표시되는 무기 프리팹이 교체됩니다. 같은 무기는 다시 장착하지 않습니다.")
    add_bullet(doc, "AI는 사용 가능한 무기를 우선 탐색하고, 무기가 없거나 목표 무기가 없으면 플레이어를 추격합니다. 장착 무기의 공격 범위 안에서는 공격합니다.")

    doc.add_heading("승패와 진행", level=2)
    add_table(
        doc,
        ["상황", "결과"],
        [
            ("AI 체력 0", "플레이어 승리, 누적 승수 +1"),
            ("플레이어 체력 0", "플레이어 패배"),
            ("제한 시간 종료 시 AI 생존", "무승부 없이 플레이어 패배"),
            ("다음 매치 시작", "현재 승수 이하의 MinWins 중 가장 높은 MatchData 선택"),
            ("재시작", "승수를 0으로 초기화하고 첫 매치 규칙으로 시작"),
        ],
        [2600, 6760],
    )

    add_callout(
        doc,
        "공격 중 이동 규칙",
        "현재 구현에서는 WeaponType.Range인 무기가 공격 중일 때 이동을 막고, 근접 무기는 이동과 공격을 동시에 허용합니다. "
        "무기 종류와 행동 규칙의 결합을 줄이려면 향후 blocksMovementWhileAttacking 같은 능력 값으로 치환할 수 있습니다.",
        fill=PALE_AMBER,
        accent="9A6700",
    )


def add_architecture(doc):
    doc.add_heading("2. 전체 설계 구조", level=1)
    add_para(doc, "구조의 핵심은 입력·판단·실행·데이터·표현을 한 클래스에 모으지 않고, 변경 이유가 다른 책임을 분리한 것입니다. Controller는 흐름을 조정하고 실제 동작은 각 기능 컴포넌트가 소유합니다.")

    doc.add_heading("설계 구조도", level=2)
    add_layer_box(doc, "표현 계층 (UI)", "MatchUI · MatchHUDUI · HealthBarUI · ResultUI", PALE_BLUE)
    add_arrow(doc, "이벤트 구독·읽기 전용 상태 조회")
    add_layer_box(doc, "흐름 조정 계층", "MatchManager · PlayerController · AIController", "E8F2F8")
    add_arrow(doc, "명령 전달·초기화")
    add_layer_box(doc, "게임플레이 기능 계층", "InputReader · Movement · Aim · Combat · Health · WeaponHolder/View · AIBrain · WeaponRuntime", PALE_GREEN)
    add_arrow(doc, "데이터 참조·비동기 리소스 요청")
    add_layer_box(doc, "데이터·리소스 계층", "Excel/SO 데이터 · Addressables · 프리팹 · 씬의 스폰 포인트·NavMesh", PALE_AMBER)
    add_para(doc, "의존 방향은 위에서 아래입니다. 하위 기능은 UI나 매치 화면을 직접 알지 않으며, UI는 전투 로직을 실행하지 않습니다.", size=9.5, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=5)

    doc.add_heading("런타임 매치 흐름", level=2)
    for number, text in enumerate((
        "시작 버튼 → MatchManager.StartGame()이 현재 승수에 맞는 MatchData를 선택합니다.",
        "MatchData가 참조하는 AIData와 PlayerData 주소로 플레이어와 AI를 순차 생성합니다.",
        "생성된 인스턴스에 데이터·타깃·WeaponSpawner를 Initialize()로 주입합니다.",
        "WeaponSpawner가 드롭 테이블과 생성 주기를 적용하고 MatchStarted 이벤트로 HUD를 연결합니다.",
        "플레이어 입력과 AI 판단이 각 Controller를 거쳐 Movement·Aim·Combat에 전달됩니다.",
        "사망 또는 시간 종료 → 입력·AI 제어 정지 → 무기 제거 → 캐릭터 Addressables 해제 → MatchEnded 이벤트로 결과 UI 표시.",
    ), start=1):
        add_number(doc, number, text)

    add_callout(
        doc,
        "비동기 시작 순서를 순차화한 이유",
        "AI 초기화에는 플레이어 Transform이 필요하고, HUD와 무기 스폰은 두 캐릭터 초기화가 끝난 뒤 시작해야 합니다. "
        "따라서 플레이어 → AI → 스포너·이벤트 순으로 완료 지점을 명확히 두었습니다.",
    )


def add_roles(doc):
    doc.add_heading("3. 역할 분리와 책임", level=1)
    add_table(
        doc,
        ["구분", "대표 클래스", "책임", "책임 밖"],
        [
            ("Manager", "MatchManager", "매치 전체 생명주기·승수·승패·생성/해제", "이동 계산, 무기 판정, UI 렌더링"),
            ("Controller", "PlayerController, AIController", "한 캐릭터의 입력/판단 결과를 기능에 전달", "기능 내부 계산과 데이터 원본 소유"),
            ("Brain", "AIBrain", "상황을 AIState와 CurrentTarget으로 결정", "NavMesh 이동, 회전, 공격 실행"),
            ("Component", "Movement, Aim, Combat, Health", "한 가지 런타임 기능과 상태 소유", "매치 진행, 화면 전환"),
            ("Runtime", "WeaponRuntime 계열", "장착된 무기의 공격 방식 구현", "누가 장착할지·드롭 확률 결정"),
            ("Data", "PlayerData 등 SO", "밸런스·주소·참조 데이터 제공", "매 프레임 행동 수행"),
            ("View/UI", "WeaponView, MatchUI 계열", "표현과 사용자 피드백", "게임 규칙의 최종 판단"),
        ],
        [1150, 2050, 3330, 2830],
        font_size=8.8,
    )

    doc.add_heading("플레이어 책임 분리", level=2)
    add_table(
        doc,
        ["클래스", "역할"],
        [
            ("InputReader", "Unity Input System의 Move·Aim·Attack·Dash를 읽고 입력 상태만 보관합니다."),
            ("PlayerController", "Aim → Attack → Dash → Move 순서로 기능을 조정하고, UI에 필요한 읽기 전용 상태를 제공합니다."),
            ("PlayerMovement", "CharacterController 이동, 중력, 대시·쿨다운, 넉백을 처리합니다."),
            ("PlayerAim", "마우스 스크린 좌표를 수평 월드 지점으로 변환해 캐릭터 회전을 적용합니다."),
            ("CharacterCombat", "현재 무기의 공격 가능 여부·쿨다운·공격 중 이동 차단 여부를 관리합니다."),
        ],
        [2350, 7010],
    )

    add_callout(
        doc,
        "PlayerController가 모든 기능을 직접 구현하지 않는 이유",
        "플레이어 입력을 전투·이동 코드 안에 직접 넣으면 AI가 같은 전투·체력·무기 기능을 재사용하기 어렵습니다. "
        "Controller는 ‘언제 호출할지’를 조정하고, 기능 컴포넌트는 ‘어떻게 동작할지’를 담당하므로 입력 변경이 전투 구현을 흔들지 않습니다.",
        fill=PALE_GREEN,
        accent=DARK_BLUE,
    )


def add_character_ai_weapon(doc):
    doc.add_heading("4. 캐릭터·AI·무기 설계", level=1)

    doc.add_heading("공통 캐릭터 기능", level=2)
    add_table(
        doc,
        ["클래스", "핵심 책임"],
        [
            ("CharacterHealthSystem", "최대/현재 체력, TakeDamage(), HealthChanged(current,max), Died 이벤트"),
            ("CharacterDamageReceiver", "DamageInfo 수신 후 체력 감소, IKnockbackReceiver에 넉백 전달"),
            ("CharacterWeaponHolder", "현재 WeaponData 소유, 장착 가능 여부·교체·WeaponChanged 이벤트"),
            ("CharacterWeaponView", "WeaponChanged를 구독해 손 소켓에 무기 프리팹 생성, WeaponRuntime 초기화·이전 무기 해제"),
            ("CharacterAnimation", "이동 속도로 Idle/Walk 전환, 사망 트리거 처리"),
        ],
        [2500, 6860],
    )

    doc.add_heading("AI 구조", level=2)
    add_para(doc, "현재 AI는 한 마리를 기준으로 하며, 복잡한 행동 트리가 필요한 수준이 아니므로 enum 기반 상태 판단을 사용합니다. AIBrain은 일반 C# 클래스로 판단만 수행하고, AIController가 실제 기능을 호출합니다.")
    add_table(
        doc,
        ["상태", "진입 조건·동작"],
        [
            ("Idle", "플레이어 타깃이 없을 때 정지"),
            ("SeekWeapon", "장착 가능한 활성 무기 픽업이 있으면 가장 가까운 픽업으로 이동"),
            ("Chase", "무기가 없거나 공격 범위 밖이면 플레이어 추격"),
            ("Engage", "무기 Range 안에서 플레이어를 바라보고 공격; 원거리 공격 중에는 정지"),
            ("Dead", "판단과 이동 중단"),
        ],
        [1800, 7560],
    )
    add_para(doc, "AIMovement는 NavMeshAgent로 경로를 탐색하고, 넉백은 agent.Move()에 감쇠 벡터를 적용합니다. 향후 맵에 장애물이 추가되어도 베이크된 NavMesh 경로를 사용할 수 있다는 점이 선택 근거입니다.", size=10.2)

    doc.add_heading("무기 런타임 구조", level=2)
    add_table(
        doc,
        ["타입", "처리 방식"],
        [
            ("WeaponRuntime", "WeaponData·Owner·공격 상태를 보유하는 추상 기반 클래스"),
            ("HitBoxWeaponRuntime", "공격 시간 동안 Collider를 켜고, 한 번의 공격에서 같은 대상 중복 타격을 HashSet으로 방지"),
            ("ProjectileWeaponRuntime", "ProjectileData 주소로 투사체 프리팹을 생성하고 방향·소유자·데미지 데이터를 전달"),
            ("Projectile", "투사체 공통 초기화, 소유자 충돌 제외, 수명 종료와 Addressables 해제"),
            ("StraightProjectile", "Rigidbody 속도로 직선 이동하고 대상 또는 비 Trigger 장애물 충돌 시 해제"),
        ],
        [2450, 6910],
    )

    add_callout(
        doc,
        "Rigidbody를 투사체에 둔 이유",
        "플레이어 이동은 직접 제어가 우선이므로 CharacterController가 적합하고, 투사체는 물리 업데이트·Trigger 충돌이 핵심입니다. "
        "필요한 객체에만 Rigidbody를 두어 플레이어·AI의 이동 안정성과 투사체 충돌 검출을 분리했습니다.",
    )


def add_data(doc):
    doc.add_heading("5. 데이터 구조와 Excel → SO 도구", level=1)
    add_para(doc, "밸런스와 리소스 주소를 Excel에서 관리하고, 에디터 도구가 ScriptableObject로 변환합니다. 런타임은 Excel을 직접 읽지 않고 생성된 SO만 참조합니다.")

    doc.add_heading("테이블 참조 관계", level=2)
    add_table(
        doc,
        ["상위 데이터", "참조", "용도"],
        [
            ("MatchData", "AIData", "해당 승수 구간에서 생성할 AI"),
            ("MatchData", "MatchDropData", "해당 매치의 무기 드롭 가중치 목록"),
            ("AIData", "AIBehaviorData", "반응 시간 등 AI 판단 성향"),
            ("MatchDropData.Entry", "WeaponData", "드롭 대상 무기와 가중치"),
            ("WeaponData", "ProjectileData (선택)", "원거리 무기의 투사체 속성·프리팹 주소"),
            ("PlayerData", "독립 데이터", "플레이어 체력·속도·대시·프리팹 주소"),
        ],
        [2300, 2600, 4460],
    )
    add_para(doc, "관계 요약: MatchData → AIData → AIBehaviorData / MatchData → MatchDropData → WeaponData → ProjectileData", size=10, bold_prefix="관계 요약:", color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

    doc.add_heading("2-Pass 변환 파이프라인", level=2)
    add_layer_box(doc, "1. 읽기", "ExcelReader: ExcelDataReader로 워크북을 열고 DataTable로 변환", PALE_BLUE)
    add_arrow(doc, "스키마·ID·값 검증")
    add_layer_box(doc, "2. 기본 필드 적용 (Import)", "DataTableValidator + ExcelValueConverter + DataImporter: 문자열·숫자 등 기본 필드로 SO 생성/갱신", PALE_GREEN)
    add_arrow(doc, "대상 SO가 모두 존재한 뒤")
    add_layer_box(doc, "3. 참조 연결 (Resolve)", "DataReferenceValidator + DataImporter: ID로 대상 SO를 찾아 일반 필드·List 요소 참조 적용", PALE_AMBER)
    add_arrow(doc, "최상위 작업 종료 시 1회")
    add_layer_box(doc, "4. 저장", "DataImporterMenu: 전체 성공 여부를 모아 AssetDatabase.SaveAssets() 호출", "F1F3F5")

    doc.add_heading("도구별 책임", level=2)
    add_table(
        doc,
        ["클래스", "역할"],
        [
            ("ExcelReader", "파일 스트림과 reader를 using으로 닫고 DataTable 반환"),
            ("DataTableValidator", "일반 테이블의 빈/중복 ID, 리스트 테이블 ID, 헤더·기본 값 검증"),
            ("ExcelValueConverter", "셀 문자열을 대상 FieldInfo 타입으로 변환"),
            ("DataReferenceValidator", "일반·List의 SO 참조가 모두 유효한지 먼저 검증하고 실패 시 부분 참조 적용 방지"),
            ("DataImporter", "검증된 값을 SO에 반영하고 기존 SO는 갱신하여 GUID와 참조 유지"),
            ("DataImporterMenu", "개별/전체 Import·Resolve 메뉴, 성공 결과 집계, 최종 저장"),
        ],
        [2450, 6910],
    )

    add_callout(
        doc,
        "Import와 Resolve를 분리한 이유",
        "Excel의 참조 대상 SO가 아직 만들어지지 않은 순서 문제를 제거하기 위해서입니다. 1-Pass에서 모든 SO의 기본 형태를 확보하고, "
        "2-Pass에서 ID 참조를 연결하므로 테이블 순서에 대한 결합이 줄고 참조 오류를 명확하게 보고할 수 있습니다.",
        fill=PALE_GREEN,
        accent=DARK_BLUE,
    )

    doc.add_heading("데이터 변경 절차", level=2)
    for number, text in enumerate((
        "Excel의 필드 헤더와 대상 SO 직렬화 필드 이름을 일치시킵니다.",
        "주소·ID·가중치·수치 값을 수정하고 파일을 저장합니다.",
        "한 테이블만 바뀌었다면 Tools/Excel Import/Individual 메뉴를 사용합니다.",
        "여러 참조 테이블이 함께 바뀌었다면 0. Import And Resolve (전체)를 실행합니다.",
        "Console의 검증 오류가 없는지 확인한 뒤 생성된 SO 참조와 실제 플레이를 점검합니다.",
    ), start=1):
        add_number(doc, number, text)


def add_addressables(doc):
    doc.add_heading("6. Addressables 운용", level=1)
    add_para(doc, "Addressables는 매치마다 달라지거나 데이터 주소로 선택되는 프리팹을 런타임에 생성하는 데 사용합니다. UI에 고정된 이미지나 프리팹 내부 종속 에셋까지 무조건 개별 주소화하지 않습니다.")

    doc.add_heading("현재 등록·사용 자원", level=2)
    add_table(
        doc,
        ["분류", "주소 예", "실제 사용 지점"],
        [
            ("플레이어", "Characters/Player", "MatchManager가 PlayerData 주소로 생성"),
            ("AI", "Characters/AI_Easy, AI_Normal, AI_Hard", "MatchData→AIData 주소로 생성"),
            ("손 무기", "Weapons/Sword, Hammer, Shield, Bow", "CharacterWeaponView가 장착 변경 시 생성"),
            ("무기 픽업", "Pickups/WeaponPickup", "WeaponSpawner가 AssetReferenceGameObject로 생성"),
            ("투사체", "Projectiles/Arrow", "ProjectileWeaponRuntime이 ProjectileData 주소로 생성"),
        ],
        [1700, 3300, 4360],
        font_size=8.9,
    )

    doc.add_heading("비동기 수명주기", level=2)
    add_table(
        doc,
        ["단계", "처리"],
        [
            ("요청", "SO에 저장된 주소 또는 AssetReference로 InstantiateAsync 호출"),
            ("완료 검증", "Status, 컴포넌트 존재, Initialize() 성공 여부 확인"),
            ("경합 방지", "matchVersion·spawnVersion·view version으로 이전 비동기 결과를 폐기"),
            ("정상 해제", "생성 인스턴스는 Addressables.ReleaseInstance()로 반환"),
            ("실패 해제", "실패 handle은 Addressables.Release(handle) 처리"),
        ],
        [1900, 7460],
    )

    add_callout(
        doc,
        "버전 검사의 목적",
        "비동기 로드 완료 전에 매치가 종료되거나 다른 무기로 교체될 수 있습니다. 요청 시점의 버전과 완료 시점의 버전이 다르면 "
        "늦게 도착한 결과를 즉시 해제하여 이전 매치·이전 무기가 현재 상태를 덮어쓰지 못하게 합니다.",
    )

    doc.add_heading("맵 운용 정책", level=2)
    add_callout(
        doc,
        "[현재] 씬 고정 맵",
        "현재 요구사항은 맵 1개이며 Match_1 프리팹 인스턴스, Player/AISpawnPoint, WeaponSpawner와 무기 스폰 지점, "
        "베이크된 NavMeshSurface를 플레이 씬에서 사용합니다. 따라서 MatchManager가 맵을 Addressables로 로드하지 않습니다.",
        fill=PALE_GREEN,
        accent=DARK_BLUE,
    )
    add_callout(
        doc,
        "[확장] 다중 맵",
        "맵이 추가되면 MatchData에 맵 주소를 추가하고 MapLoader 또는 MapManager가 프리팹을 생성하도록 확장합니다. "
        "맵 프리팹의 MapContext가 Player/AISpawnPoint, WeaponSpawnPoints, NavMeshSurface를 외부에 제공하면 씬 참조 단절을 막을 수 있습니다.",
        fill=PALE_AMBER,
        accent="9A6700",
    )


def add_design_decisions(doc):
    doc.add_heading("7. 주요 설계 선택과 근거", level=1)
    add_table(
        doc,
        ["선택", "적용", "선택 근거"],
        [
            ("플레이어 이동", "CharacterController", "정교한 연쇄 물리보다 직접 제어·충돌·넉백 재현성이 중요"),
            ("AI 이동", "NavMeshAgent", "현재 맵과 향후 장애물 맵에서 경로 탐색을 재사용하고 난이도별 판단과 이동을 분리"),
            ("AI 판단", "enum 상태 + AIBrain", "현재 상태 수와 AI 1마리 규모에 맞고, 완전한 State 패턴·Behavior Tree의 객체/설정 비용을 피함"),
            ("공통 전투", "기능 컴포넌트 공유", "플레이어 입력과 AI 판단이 달라도 Health·Combat·Weapon 기능은 동일하게 재사용"),
            ("데미지 전달", "DamageInfo 구조체", "한 번의 타격에 함께 이동하는 값 묶음이며 독립 생명주기·상속이 필요하지 않음"),
            ("UI 갱신", "이벤트 + 읽기 상태", "체력·매치 전환은 이벤트, 연속 쿨다운은 HUD가 읽어 갱신해 불필요한 결합을 줄임"),
            ("데이터", "Excel + SO", "기획 데이터 편집성과 Unity 런타임 참조 안정성을 동시에 확보"),
            ("리소스", "선택적 Addressables", "데이터로 바뀌는 프리팹은 주소 로드, 항상 함께 필요한 종속 에셋은 프리팹 참조 유지"),
        ],
        [1800, 2200, 5360],
        font_size=8.7,
    )

    doc.add_heading("상태머신과 Behavior Tree를 지금 사용하지 않은 이유", level=2)
    add_para(doc, "플레이어는 이동·에임·공격이 동시에 일어나므로 하나의 단일 상태로 제한하면 조합 상태가 급격히 늘어납니다. 반면 AI는 SeekWeapon·Chase·Engage처럼 우선 행동을 하나 선택해야 하므로 상태 구분이 유효합니다. 현재 판단은 선형 우선순위로 충분하고, 조건 조합·병렬 행동·서브트리가 크게 늘 때 Behavior Tree 도입을 검토합니다.")

    doc.add_heading("Manager를 제한적으로 사용한 이유", level=2)
    add_para(doc, "전역 접근이 필요한 클래스마다 Manager를 만들지 않았습니다. MatchManager만 매치 전체 생명주기를 소유하고, WeaponSpawner는 무기 생성이라는 장면 내 책임에 한정합니다. Addressables 서비스나 오브젝트 풀은 중복 로딩 정책·캐시·풀링 요구가 실제로 생길 때 분리하는 편이 현재 규모에서 책임과 호출 경로를 더 명확하게 유지합니다.")


def add_handoff(doc):
    doc.add_heading("8. 인계 및 운영 가이드", level=1)

    doc.add_heading("씬 필수 구성", level=2)
    add_table(
        doc,
        ["항목", "필수 연결"],
        [
            ("MatchManager", "MatchData 목록, PlayerData, WeaponSpawner, Player/AISpawnPoint, deathAnimationDuration"),
            ("WeaponSpawner", "WeaponPickup AssetReference, 무기 SpawnPoints"),
            ("플레이 맵", "콜라이더, 베이크된 NavMeshSurface, 이동 가능 영역"),
            ("UI", "GameStart/Match/Result Canvas와 MatchUI·HUD·Result 연결"),
            ("캐릭터 프리팹", "Controller, Movement/Aim, Health/DamageReceiver, Combat, Holder/View, Animation"),
            ("무기 프리팹", "WeaponRuntime 파생 컴포넌트, HitBox 또는 Projectile spawn point"),
        ],
        [2350, 7010],
    )

    doc.add_heading("새 콘텐츠 추가 절차", level=2)
    add_table(
        doc,
        ["추가 대상", "작업 순서"],
        [
            ("AI", "AI/Behavior Excel 행 추가 → 프리팹 Addressable 등록 → 주소 입력 → Import/Resolve → MatchData 참조 확인"),
            ("근접 무기", "WeaponData 행·프리팹·HitBoxWeaponRuntime 구성 → Addressable 등록 → DropList 가중치 연결"),
            ("원거리 무기", "ProjectileData·투사체 프리팹 추가 → ProjectileWeaponRuntime 무기 연결 → 두 주소 등록 → Import/Resolve"),
            ("매치 단계", "MatchData에 MinWins·AI·DropList·시간 설정 → MatchManager 목록에 SO 추가"),
            ("맵", "현재는 씬 교체/수정; 다중 맵 요구 시 MapContext 기반 Addressable 로더부터 도입"),
        ],
        [2000, 7360],
        font_size=9.0,
    )

    doc.add_heading("제출·인계 전 점검", level=2)
    add_table(
        doc,
        ["우선도", "점검 항목", "판정 기준"],
        [
            ("필수", "Addressables Content Build", "Player/AI/무기/픽업/투사체가 에디터와 빌드에서 모두 생성·해제"),
            ("필수", "전체 Excel Import + Resolve", "빈·중복 ID, 타입 변환, 참조 오류 없이 SO 갱신"),
            ("필수", "게임 루프", "시작→진행→승/패→다음/재시작이 반복 가능"),
            ("필수", "씬 참조", "MatchManager·WeaponSpawner·스폰 포인트·UI·NavMeshSurface 연결"),
            ("정리", "Matches/Match_1 Addressable 등록", "현재 씬 고정 정책이면 사용되지 않는 등록을 제거"),
            ("정리", "StraightProjectile Rigidbody", "projectileRigidbody 필드를 실제 사용하거나 필드를 제거하여 중복 접근 제거"),
            ("선택", "WeaponType 의존", "무기 행동 조합이 늘면 이동 차단을 capability bool로 이전"),
            ("제출", "외부 무료 에셋 출처", "프로젝트 내 별도 메모 또는 제출 문서에 출처 기록"),
        ],
        [1000, 3450, 4910],
        font_size=8.6,
    )

    add_callout(
        doc,
        "현재 알려진 코드 정리 항목",
        "StraightProjectile은 projectileRigidbody 필드를 선언했지만 TryShoot()에서 GetComponent<Rigidbody>()를 다시 호출합니다. "
        "기능 오류는 아니지만 필드를 캐시할 목적이었다면 해당 필드를 사용하고, 아니라면 선언을 제거해야 의도가 명확합니다. "
        "또한 Matches/Match_1은 Addressables 그룹에 남아 있으나 현재 MatchManager에서 로드하지 않으므로 씬 고정 정책 확정 시 등록을 정리합니다.",
        fill=PALE_RED,
        accent="A13D3D",
    )


def add_test_matrix(doc):
    doc.add_heading("9. 기능 검증 매트릭스", level=1)
    add_table(
        doc,
        ["영역", "검증 시나리오", "기대 결과"],
        [
            ("입력", "WASD·마우스·공격·대시", "이동/에임 독립, 1회성 입력 소비, 쿨다운 준수"),
            ("전투", "근접·활 공격", "HitBox/투사체 데미지·넉백, 같은 공격 중 중복 타격 방지"),
            ("원거리 이동", "활 공격 중 이동 입력", "공격 지속 중 정지, 종료 후 이동 복구"),
            ("무기 교체", "서로 다른 무기/같은 무기 픽업", "다른 무기 교체, 같은 무기 무시, 이전 View 해제"),
            ("AI", "무기 있음/없음/범위 내외", "SeekWeapon→Chase→Engage 전환, 플레이어 방향 갱신"),
            ("스폰", "복수 스폰 주기·실패", "점유 위치 중복 방지, 실패 예약 복구, 종료 시 모두 해제"),
            ("승패", "AI 사망/플레이어 사망/시간 종료", "승/패·승수·결과 UI가 규칙대로 처리"),
            ("데이터", "개별/전체 Import", "기존 GUID 유지, 잘못된 ID·참조 차단, 성공 시 한 번 저장"),
            ("Addressables", "매치 반복·빠른 교체", "오래된 완료 콜백이 현재 상태를 덮지 않고 인스턴스 누수 없음"),
        ],
        [1200, 3900, 4260],
        font_size=8.5,
    )

    doc.add_heading("리뷰 인터뷰에서 설명할 핵심", level=2)
    add_bullet(doc, "Excel은 편집 원본, SO는 Unity 런타임 데이터입니다. 2-Pass로 생성과 참조 연결을 분리해 순서 문제와 부분 적용을 줄였습니다.", keep_next=True)
    add_bullet(doc, "Addressables는 ‘모든 에셋’이 아니라 런타임에 데이터로 선택·교체되는 프리팹에 적용했습니다. 생성과 해제 책임은 요청한 클래스가 가집니다.", keep_next=True)
    add_bullet(doc, "플레이어와 AI의 차이는 입력/판단이며, 체력·전투·무기 기능은 공통 컴포넌트로 재사용합니다.", keep_next=True)
    add_bullet(doc, "현재 요구사항에 맞게 enum AI 상태와 씬 고정 맵을 선택했고, 복잡도가 증가할 때의 Behavior Tree·MapContext 확장 지점을 남겼습니다.", keep_next=True)
    add_bullet(doc, "예외 처리는 가능한 모든 상황을 막는 것이 아니라, 비동기 로드·데이터 무결성·인스턴스 해제처럼 실패 비용이 큰 경계에 집중했습니다.", keep_next=True)

    add_callout(
        doc,
        "인계 결론",
        "프로젝트는 데이터 편집(Excel), Unity 데이터(SO), 리소스 로드(Addressables), 매치 흐름(MatchManager), 캐릭터 기능(Component), "
        "표현(UI)을 분리한 구조입니다. 새 AI·무기·매치 데이터는 기존 파이프라인으로 추가할 수 있으며, 맵·풀링·사운드처럼 현재 요구 밖의 시스템은 필요 시 확장하도록 경계를 남겼습니다.",
        fill=PALE_BLUE,
    )


def add_appendix(doc):
    doc.add_heading("부록. 주요 파일 위치", level=1)
    add_table(
        doc,
        ["영역", "경로"],
        [
            ("매치", "Assets/3.Scripts/Match/MatchManager.cs"),
            ("플레이어", "Assets/3.Scripts/2.Player/"),
            ("AI", "Assets/3.Scripts/3.AI/"),
            ("공통 캐릭터", "Assets/3.Scripts/1.Character/"),
            ("무기", "Assets/3.Scripts/Weapon/"),
            ("UI", "Assets/3.Scripts/UI/"),
            ("런타임 데이터", "Assets/3.Scripts/RunTime/"),
            ("Excel 도구", "Assets/Editor/Scirpts/ExcelTool/"),
            ("Excel 원본", "Assets/Editor/ExcelData/"),
            ("생성된 SO", "Assets/8.GameData/"),
            ("Addressables 설정", "Assets/AddressableAssetsData/"),
            ("플레이 씬", "Assets/1.Scenes/PlayerMatchScene.unity"),
        ],
        [2300, 7060],
        font_size=9.1,
    )
    add_para(doc, "주의: 폴더명 ‘Scirpts’는 현재 저장소의 실제 경로 표기를 그대로 사용했습니다.", size=9.5, color=MUTED, italic=True)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_contents(doc)
    add_game_overview(doc)
    add_architecture(doc)
    add_roles(doc)
    add_character_ai_weapon(doc)
    add_data(doc)
    add_addressables(doc)
    add_design_decisions(doc)
    add_handoff(doc)
    add_test_matrix(doc)
    add_appendix(doc)

    props = doc.core_properties
    props.title = "SHINZI Games 인턴십 과제 프로젝트 인계 및 설계 문서"
    props.subject = "Unity 탑다운 액션 게임 구조, Excel→SO, Addressables 인계 가이드"
    props.keywords = "Unity, Excel, ScriptableObject, Addressables, Top-down, Handoff"
    props.comments = "현재 저장소 구현 기준"

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
